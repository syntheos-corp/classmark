#!/usr/bin/env python3
"""
Classification Marking Scanner - Desktop Edition
Scans documents (PDF, DOCX, TXT, JSON) for US Government security classification markings.

Features:
- Multi-format support (PDF, DOCX, TXT, JSON)
- Pattern matching with CAPCO standard markings
- Fuzzy matching for variations
- Local LLM verification via Ollama (optional)
- High recall configuration
- Batch processing with parallel workers
- Detailed reporting with confidence scoring

Author: Claude Code
Date: 2025-01-05
Version: 2.1 Desktop Edition
"""

import os
import re
import json
import argparse
import sys
import hashlib
import shutil
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Set
from dataclasses import dataclass, asdict
from enum import Enum
from collections import defaultdict
import warnings
warnings.filterwarnings('ignore')

# Configuration
from config import MAX_FILE_SIZE_BYTES, MAX_FILE_SIZE_MB

# Progress tracking
try:
    from tqdm import tqdm
    HAS_TQDM = True
except ImportError:
    HAS_TQDM = False

# Parallel processing
from concurrent.futures import ProcessPoolExecutor, as_completed
import multiprocessing

# PDF processing
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# DOCX processing
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import docx2python
    HAS_DOCX2PYTHON = True
except ImportError:
    HAS_DOCX2PYTHON = False

# Fuzzy matching
try:
    from rapidfuzz import fuzz, process
    HAS_RAPIDFUZZ = True
except ImportError:
    try:
        from fuzzywuzzy import fuzz, process
        HAS_RAPIDFUZZ = True
    except ImportError:
        HAS_RAPIDFUZZ = False

# Local LLM
try:
    import ollama
    HAS_OLLAMA = True
except ImportError:
    HAS_OLLAMA = False

# OCR
try:
    from PIL import Image
    import pytesseract
    from pdf2image import convert_from_path
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


class ClassificationLevel(Enum):
    """US Government classification levels per CAPCO standards"""
    TOP_SECRET = "TOP SECRET"
    SECRET = "SECRET"
    CONFIDENTIAL = "CONFIDENTIAL"
    CUI = "CONTROLLED UNCLASSIFIED INFORMATION"
    UNCLASSIFIED = "UNCLASSIFIED"
    UNKNOWN = "UNKNOWN"

    def __lt__(self, other):
        """Allow sorting by classification level"""
        order = [
            ClassificationLevel.TOP_SECRET,
            ClassificationLevel.SECRET,
            ClassificationLevel.CONFIDENTIAL,
            ClassificationLevel.CUI,
            ClassificationLevel.UNKNOWN,
            ClassificationLevel.UNCLASSIFIED,
        ]
        return order.index(self) < order.index(other)


@dataclass
class Match:
    """Individual classification marking match"""
    text: str
    position: int
    line_number: int
    confidence: float
    context: str
    match_type: str  # 'pattern', 'fuzzy', 'structural'
    location: str  # 'header', 'footer', 'body', 'metadata'


@dataclass
class DetectionResult:
    """Result of classification marking detection for a single file"""
    file_path: str
    file_hash: str
    has_markings: bool
    classification_level: ClassificationLevel
    matches: List[Match]
    overall_confidence: float
    ocr_used: bool
    llm_verified: bool
    processing_time: float
    file_size: int


class ClassificationPatterns:
    """CAPCO-compliant classification marking patterns"""

    # Classification levels
    TOP_SECRET_PATTERNS = [
        r'(?:^|\n)\s*(?:\/\/|//)?\s*TOP\s*SECRET(?:\s*\/\/[A-Z\s\/]+)?',
        r'\(TS(?:\/\/[A-Z]+)?\)',
        r'TOP\s*SECRET\s*\/\/\s*(?:SCI|NOFORN|ORCON|IMCON|RELIDO|PROPIN|FISA)',
    ]

    SECRET_PATTERNS = [
        r'(?:^|\n)\s*(?:\/\/|//)?\s*SECRET(?:\s*\/\/[A-Z\s\/]+)?(?!\s*(?:SERVICE|AGENT|SANTA))',
        r'\(S(?:\/\/[A-Z]+)?\)(?!\w)',  # (S) but not part of other words
        r'SECRET\s*\/\/\s*(?:NOFORN|ORCON|RELIDO|REL)',
    ]

    CONFIDENTIAL_PATTERNS = [
        r'(?:^|\n)\s*(?:\/\/|//)?\s*CONFIDENTIAL(?:\s*\/\/[A-Z\s\/]+)?',
        r'\(C(?:\/\/[A-Z]+)?\)(?!\w)',
        r'CONFIDENTIAL\s*\/\/\s*(?:NOFORN|RELIDO)',
    ]

    CUI_PATTERNS = [
        r'(?:^|\n)\s*(?:CONTROLLED\s+UNCLASSIFIED\s+INFORMATION|CUI)(?:\s*\/\/[A-Z\s\/]+)?',
        r'\(CUI\)',
        r'CUI\s*\/\/\s*[A-Z]+',
    ]

    # Legacy CUI markings (per DoD transition guidance, treat as CUI equivalent)
    LEGACY_CUI_PATTERNS = [
        r'\bFOUO\b',  # For Official Use Only
        r'\bFOR\s+OFFICIAL\s+USE\s+ONLY\b',
        r'\bSBU\b',  # Sensitive But Unclassified
        r'\bSENSITIVE\s+BUT\s+UNCLASSIFIED\b',
        r'\bLES\b',  # Law Enforcement Sensitive
        r'\bLAW\s+ENFORCEMENT\s+SENSITIVE\b',
        r'\bLIMDIS\b',  # Limited Distribution
        r'\bLIMITED\s+DISTRIBUTION\b',
    ]

    # Portion markings at start of paragraphs
    PORTION_MARK_PATTERNS = [
        r'(?:^|\n)\s*\((TS|S|C|U)(?:\/\/[A-Z]+)?\)\s+[A-Z][a-z]{2,}',  # Require at least 3-letter word
    ]

    # Classification authority blocks
    AUTHORITY_BLOCK_PATTERNS = [
        r'CLASSIFIED\s+BY:?\s*[^\n]+',
        r'(?:^|\n)\s*DERIVED\s+FROM:?\s*[^\n]+',  # Must be at line start
        r'DECLASSIFY\s+ON:?\s*[^\n]+',
        r'REASON:?\s*1\.\d\([a-h]\)',
        r'CLASSIFICATION\s+AUTHORITY:',
    ]

    # Banner patterns (repeated markings)
    BANNER_PATTERNS = [
        # Fixed ReDoS vulnerability by limiting match window to 500 chars and using lazy quantifier
        r'(?:TOP\s+SECRET|SECRET|CONFIDENTIAL)\s*(?:\/\/[A-Z\s]+)?\s*\n.{0,500}?\n\s*(?:TOP\s+SECRET|SECRET|CONFIDENTIAL)',
    ]

    # Control markings
    CONTROL_MARKINGS = [
        r'NOFORN',  # Not releasable to foreign nationals
        r'ORCON',   # Originator controlled
        r'IMCON',   # Controlled imagery
        r'RELIDO',  # Releasable by information disclosure official
        r'PROPIN',  # Caution-proprietary information involved
        r'FISA',    # Foreign Intelligence Surveillance Act
        r'SCI',     # Sensitive Compartmented Information
    ]

    # Declassification indicators (documents that were or are classified)
    DECLASSIFICATION_PATTERNS = [
        r'\bDECLASSIFIED\b',  # Declassification stamps
        r'RELEASE\s+AS\s+SANITIZED',  # CIA historical releases
        r'APPROVED\s+FOR\s+RELEASE',  # FOIA releases
        r'CIA\s+HISTORICAL\s+REVIEW',  # CIA historical review program
        r'FOIA\s+CASE\s+NUMBER',  # Freedom of Information Act cases
        r'AUTOMATICALLY\s+DECLASSIFIED',  # Automatic declassification
        r'DOWNGRADED\s+TO',  # Downgraded classifications
        r'UPGRADED\s+TO',  # Upgraded classifications
        r'REGRADED\s+TO',  # Regraded classifications
        r'GROUP\s+1\s+-\s+EXCLUDED',  # Pre-1990s exclusion markings
        r'SANITIZED\s+COPY',  # Sanitized document copies
        r'\bREDACTED\b',  # Redacted content indicator
    ]

    @classmethod
    def get_all_patterns(cls) -> List[Tuple[str, str, re.Pattern]]:
        """
        Get all compiled patterns with their categories and names
        Returns: List of (category, name, compiled_pattern)
        """
        patterns = []

        for pattern_str in cls.TOP_SECRET_PATTERNS:
            patterns.append(('level', 'TOP_SECRET',
                           re.compile(pattern_str, re.IGNORECASE | re.MULTILINE)))

        for pattern_str in cls.SECRET_PATTERNS:
            patterns.append(('level', 'SECRET',
                           re.compile(pattern_str, re.IGNORECASE | re.MULTILINE)))

        for pattern_str in cls.CONFIDENTIAL_PATTERNS:
            patterns.append(('level', 'CONFIDENTIAL',
                           re.compile(pattern_str, re.IGNORECASE | re.MULTILINE)))

        for pattern_str in cls.CUI_PATTERNS:
            patterns.append(('level', 'CUI',
                           re.compile(pattern_str, re.IGNORECASE | re.MULTILINE)))

        for pattern_str in cls.LEGACY_CUI_PATTERNS:
            patterns.append(('level', 'LEGACY_CUI',
                           re.compile(pattern_str, re.IGNORECASE | re.MULTILINE)))

        for pattern_str in cls.PORTION_MARK_PATTERNS:
            patterns.append(('structural', 'PORTION_MARK',
                           re.compile(pattern_str, re.MULTILINE)))

        for pattern_str in cls.AUTHORITY_BLOCK_PATTERNS:
            patterns.append(('structural', 'AUTHORITY_BLOCK',
                           re.compile(pattern_str, re.IGNORECASE | re.MULTILINE)))

        for pattern_str in cls.BANNER_PATTERNS:
            patterns.append(('structural', 'BANNER',
                           re.compile(pattern_str, re.IGNORECASE | re.MULTILINE)))

        for pattern_str in cls.CONTROL_MARKINGS:
            patterns.append(('control', pattern_str,
                           re.compile(r'\b' + pattern_str + r'\b', re.IGNORECASE)))

        for pattern_str in cls.DECLASSIFICATION_PATTERNS:
            patterns.append(('declassification', 'DECLASSIFIED',
                           re.compile(pattern_str, re.IGNORECASE | re.MULTILINE)))

        return patterns


class ContextAnalyzer:
    """Analyze context around matches to reduce false positives"""

    # Indicators of OFFICIAL classification markings
    OFFICIAL_INDICATORS = [
        r'classified\s+by',
        r'derived\s+from',
        r'declassify\s+on',
        r'reason:\s*1\.\d+\([a-h]\)',
        r'classification\s+authority',
        r'marking\s+indicator',
        r'portion\s+mark',
        r'overall\s+classification',
        r'security\s+classification\s+guide',
        r'scg',  # Security Classification Guide
        r'oci',  # Original Classification Authority
    ]

    # Indicators of FALSE POSITIVES (casual usage)
    FALSE_POSITIVE_INDICATORS = [
        r'the\s+secret\s+to',
        r'keep\s+(?:it\s+)?(?:a\s+)?secret',
        r'in\s+secret',
        r'victoria\'?s?\s+secret',
        r'secret\s+sauce',
        r'secret\s+santa',
        r'secret\s+service',
        r'secret\s+agent',
        r'secret\s+ingredient',
        r'trade\s+secret',
        r'classified\s+as\s+(?:a|an)',
        r'classified\s+into',
        r'can\s+be\s+classified',
        r'are\s+classified\s+as',
        r'is\s+classified\s+as',
        r'be\s+classified\s+as',
        r'confidential\s+(?:conversation|discussion|talk|chat)',
        r'derived\s+from\s+(?:the|a|an|soviet|russian|chinese)',  # Casual "derived from" usage
        r'secret\s+(?:report|document|paper|study)',  # Reference to secret docs, not markings
        r'confidential\s+(?:report|document|paper|study|source)',
    ]

    def __init__(self):
        self.official_pattern = re.compile(
            '|'.join(self.OFFICIAL_INDICATORS),
            re.IGNORECASE
        )
        self.false_positive_pattern = re.compile(
            '|'.join(self.FALSE_POSITIVE_INDICATORS),
            re.IGNORECASE
        )

    def analyze(self, text: str, match_pos: int, window: int = 300) -> Dict:
        """Analyze context around a match"""
        start = max(0, match_pos - window)
        end = min(len(text), match_pos + window)
        context = text[start:end]

        has_official = bool(self.official_pattern.search(context))
        has_false_positive = bool(self.false_positive_pattern.search(context))

        # Check if at start of line (banners are usually at line start)
        line_start = text.rfind('\n', 0, match_pos)
        chars_from_line_start = match_pos - line_start - 1
        at_line_start = chars_from_line_start < 5

        # Check if surrounded by slashes (e.g., //SECRET//)
        has_slashes = bool(re.search(r'\/\/.*\/\/', context))

        return {
            'context': context,
            'has_official': has_official,
            'has_false_positive': has_false_positive,
            'at_line_start': at_line_start,
            'has_slashes': has_slashes,
        }


class TextExtractor:
    """Extract text from various document formats"""

    # DPI requirements per government standards
    # 200 DPI minimum for 10pt+ fonts
    # 300 DPI minimum for smaller fonts, technical drawings, signatures
    DPI_MINIMUM_STANDARD = 200
    DPI_RECOMMENDED_HIGH_QUALITY = 300

    def __init__(self):
        pass

    def validate_image_dpi(self, image, page_num: int = 0) -> Tuple[bool, str]:
        """
        Validate scanned document meets minimum DPI requirements
        per government standards (32 CFR requirements)

        Returns: (meets_standards: bool, message: str)
        """
        dpi = image.info.get('dpi', (0, 0))

        # Handle both tuple and single value
        if isinstance(dpi, tuple):
            dpi_x, dpi_y = dpi
        else:
            dpi_x = dpi_y = dpi

        # Check against minimum standards
        if dpi_x < self.DPI_MINIMUM_STANDARD or dpi_y < self.DPI_MINIMUM_STANDARD:
            return False, (
                f"Page {page_num}: DPI {dpi} below minimum {self.DPI_MINIMUM_STANDARD}. "
                f"OCR quality may be significantly reduced. "
                f"Government standards require 200 DPI for 10pt+ fonts, "
                f"300 DPI for smaller fonts/technical content."
            )

        elif dpi_x < self.DPI_RECOMMENDED_HIGH_QUALITY or dpi_y < self.DPI_RECOMMENDED_HIGH_QUALITY:
            return True, (
                f"Page {page_num}: DPI {dpi} meets minimum but below recommended {self.DPI_RECOMMENDED_HIGH_QUALITY}. "
                f"Consider rescanning for better accuracy."
            )

        else:
            return True, f"Page {page_num}: DPI {dpi} meets standards (≥{self.DPI_MINIMUM_STANDARD})."

    def calculate_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def check_file_size(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Check if file size is within acceptable limits

        Returns:
            tuple: (is_valid, error_message)
        """
        try:
            file_size = os.path.getsize(file_path)
            if file_size > MAX_FILE_SIZE_BYTES:
                size_mb = file_size / (1024 * 1024)
                return False, f"File too large ({size_mb:.1f} MB, max: {MAX_FILE_SIZE_MB} MB)"
            return True, None
        except OSError as e:
            return False, f"Cannot check file size: {str(e)}"

    def extract_from_pdf_pypdf2(self, file_path: str) -> Dict[str, str]:
        """Extract text from PDF using PyPDF2"""
        result = {'body': [], 'headers': [], 'footers': [], 'metadata': {}}

        try:
            # Check file size before processing
            size_ok, error_msg = self.check_file_size(file_path)
            if not size_ok:
                print(f"Skipping {file_path}: {error_msg}", file=sys.stderr)
                return result

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract metadata
                if pdf_reader.metadata:
                    result['metadata'] = {
                        key: str(value) for key, value in pdf_reader.metadata.items()
                    }

                # Extract text from each page with region analysis
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text()
                    if page_text:
                        # Extract page regions
                        regions = self.extract_page_regions(page_text, page_num)

                        # Add regions to result
                        if regions['header']:
                            result['headers'].append(f"[PAGE {page_num} HEADER]\n{regions['header']}")
                        if regions['body']:
                            result['body'].append(f"[PAGE {page_num}]\n{regions['body']}")
                        if regions['footer']:
                            result['footers'].append(f"[PAGE {page_num} FOOTER]\n{regions['footer']}")
        except Exception as e:
            print(f"PyPDF2 error on {file_path}: {e}", file=sys.stderr)

        return result

    def extract_page_regions(self, page_text: str, page_num: int) -> Dict[str, str]:
        """
        Extract header, body, and footer regions from page text
        Header: top 10% of lines, Footer: bottom 10% of lines
        """
        if not page_text:
            return {'header': '', 'body': '', 'footer': '', 'page': page_num}

        lines = page_text.split('\n')
        total_lines = len(lines)

        if total_lines == 0:
            return {'header': '', 'body': '', 'footer': '', 'page': page_num}

        # Calculate region boundaries (top/bottom 10%)
        header_threshold = max(1, total_lines // 10)
        footer_threshold = max(1, total_lines // 10)

        # Extract regions
        header_lines = lines[:header_threshold]
        footer_lines = lines[-footer_threshold:] if total_lines > footer_threshold else []
        body_lines = lines[header_threshold:-footer_threshold] if total_lines > header_threshold + footer_threshold else lines

        return {
            'header': '\n'.join(header_lines),
            'body': '\n'.join(body_lines),
            'footer': '\n'.join(footer_lines),
            'page': page_num
        }

    def extract_from_pdf_pdfplumber(self, file_path: str) -> Dict[str, str]:
        """Extract text from PDF using pdfplumber (better layout preservation)"""
        result = {'body': [], 'headers': [], 'footers': [], 'metadata': {}}

        try:
            # Check file size before processing
            size_ok, error_msg = self.check_file_size(file_path)
            if not size_ok:
                print(f"Skipping {file_path}: {error_msg}", file=sys.stderr)
                return result

            with pdfplumber.open(file_path) as pdf:
                # Extract metadata
                if pdf.metadata:
                    result['metadata'] = {
                        key: str(value) for key, value in pdf.metadata.items()
                    }

                # Extract text from each page with region analysis
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()
                    if page_text:
                        # Extract page regions
                        regions = self.extract_page_regions(page_text, page_num)

                        # Add regions to result
                        if regions['header']:
                            result['headers'].append(f"[PAGE {page_num} HEADER]\n{regions['header']}")
                        if regions['body']:
                            result['body'].append(f"[PAGE {page_num}]\n{regions['body']}")
                        if regions['footer']:
                            result['footers'].append(f"[PAGE {page_num} FOOTER]\n{regions['footer']}")
        except Exception as e:
            print(f"pdfplumber error on {file_path}: {e}", file=sys.stderr)

        return result

    def extract_from_pdf(self, file_path: str) -> Dict[str, str]:
        """Extract text from PDF, trying best available library"""
        if HAS_PDFPLUMBER:
            return self.extract_from_pdf_pdfplumber(file_path)
        elif HAS_PYPDF2:
            return self.extract_from_pdf_pypdf2(file_path)
        else:
            return {'body': [], 'metadata': {}}

    def pdf_needs_ocr(self, extracted_data: Dict) -> bool:
        """
        Check if a PDF needs OCR (has little or no extractable text)
        Returns True if OCR should be attempted
        """
        if not extracted_data.get('body'):
            return True

        # Check if extracted text is substantial
        total_text = ' '.join(str(page) for page in extracted_data['body'])
        word_count = len(total_text.split())

        # If less than 50 words extracted, likely a scanned PDF
        return word_count < 50

    def extract_from_pdf_with_ocr(self, file_path: str) -> Dict[str, str]:
        """Extract text from scanned PDF using OCR"""
        result = {'body': [], 'metadata': {}}

        if not HAS_OCR:
            return result

        try:
            # Check file size before processing
            size_ok, error_msg = self.check_file_size(file_path)
            if not size_ok:
                print(f"Skipping {file_path}: {error_msg}", file=sys.stderr)
                return result

            # Convert PDF pages to images
            # Using lower DPI for faster processing, can increase for production
            images = convert_from_path(file_path, dpi=200, thread_count=2)

            # OCR each page
            for page_num, image in enumerate(images, start=1):
                try:
                    # Validate DPI
                    meets_standards, dpi_message = self.validate_image_dpi(image, page_num)
                    if not meets_standards:
                        print(f"Warning: {dpi_message}", file=sys.stderr)
                        # Still attempt OCR but with warning

                    # Perform OCR
                    # --psm 1: Automatic page segmentation with OSD (Orientation and script detection)
                    text = pytesseract.image_to_string(image, config='--psm 1')

                    if text.strip():
                        result['body'].append(f"[PAGE {page_num} - OCR]\n{text}")

                except Exception as e:
                    print(f"OCR error on page {page_num} of {file_path}: {e}", file=sys.stderr)
                    continue

        except Exception as e:
            print(f"PDF to image conversion error on {file_path}: {e}", file=sys.stderr)

        return result

    def extract_from_docx(self, file_path: str) -> Dict[str, str]:
        """Extract text from DOCX including headers and footers"""
        result = {'body': [], 'headers': [], 'footers': [], 'metadata': {}}

        if not HAS_DOCX:
            return result

        try:
            # Check file size before processing
            size_ok, error_msg = self.check_file_size(file_path)
            if not size_ok:
                print(f"Skipping {file_path}: {error_msg}", file=sys.stderr)
                return result

            doc = Document(file_path)

            # Extract core properties (metadata)
            if doc.core_properties:
                props = doc.core_properties
                result['metadata'] = {
                    'author': props.author or '',
                    'title': props.title or '',
                    'subject': props.subject or '',
                    'keywords': props.keywords or '',
                    'created': str(props.created) if props.created else '',
                    'modified': str(props.modified) if props.modified else '',
                }

            # Extract body paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    result['body'].append(para.text)

            # Extract headers and footers from each section
            for section in doc.sections:
                # Headers
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            result['headers'].append(para.text)

                # Footers
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            result['footers'].append(para.text)

            # Try docx2python for enhanced extraction if available
            if HAS_DOCX2PYTHON:
                try:
                    doc2 = docx2python.docx2python(file_path)
                    # Get additional headers/footers that might have been missed
                    if doc2.header:
                        result['headers'].extend([str(h) for h in doc2.header if h])
                    if doc2.footer:
                        result['footers'].extend([str(f) for f in doc2.footer if f])
                except Exception:
                    # docx2python extraction failed, continue with what we have from python-docx
                    pass

        except Exception as e:
            print(f"DOCX error on {file_path}: {e}", file=sys.stderr)

        return result

    def extract_from_txt(self, file_path: str) -> Dict[str, str]:
        """Extract text from text file"""
        result = {'body': []}

        try:
            # Check file size before processing
            size_ok, error_msg = self.check_file_size(file_path)
            if not size_ok:
                print(f"Skipping {file_path}: {error_msg}", file=sys.stderr)
                return result

            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        result['body'].append(f.read())
                    break
                except UnicodeDecodeError:
                    continue
        except Exception as e:
            print(f"Text file error on {file_path}: {e}", file=sys.stderr)

        return result

    def extract_from_json(self, file_path: str) -> Dict[str, str]:
        """Extract text from JSON file"""
        result = {'body': []}

        try:
            # Check file size before processing
            size_ok, error_msg = self.check_file_size(file_path)
            if not size_ok:
                print(f"Skipping {file_path}: {error_msg}", file=sys.stderr)
                return result

            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Convert to formatted string to catch markings in values
            result['body'].append(json.dumps(data, indent=2))
        except Exception as e:
            print(f"JSON error on {file_path}: {e}", file=sys.stderr)

        return result

    def extract(self, file_path: str) -> Tuple[Dict[str, any], bool]:
        """
        Extract text from file based on extension
        Returns: (extracted_data_dict, ocr_used_bool)
        """
        ext = Path(file_path).suffix.lower()
        ocr_used = False

        if ext == '.pdf':
            # Try normal extraction first
            extracted = self.extract_from_pdf(file_path)

            # Check if OCR is needed (scanned PDF with little text)
            if HAS_OCR and self.pdf_needs_ocr(extracted):
                ocr_result = self.extract_from_pdf_with_ocr(file_path)
                if ocr_result.get('body'):
                    # Merge or replace with OCR results
                    if extracted.get('body'):
                        extracted['body'].extend(ocr_result['body'])
                    else:
                        extracted['body'] = ocr_result['body']
                    ocr_used = True

        elif ext in ['.docx', '.doc']:
            extracted = self.extract_from_docx(file_path)
        elif ext == '.json':
            extracted = self.extract_from_json(file_path)
        elif ext in ['.txt', '.text', '.md', '.log', '.csv']:
            extracted = self.extract_from_txt(file_path)
        else:
            # Try as text file by default
            extracted = self.extract_from_txt(file_path)

        return extracted, ocr_used


class FuzzyMatcher:
    """Fuzzy matching for handling typos and variations"""

    # Known classification terms for fuzzy matching
    CLASSIFICATION_TERMS = [
        'TOP SECRET', 'SECRET', 'CONFIDENTIAL', 'UNCLASSIFIED',
        'CUI', 'CONTROLLED UNCLASSIFIED INFORMATION',
        'NOFORN', 'ORCON', 'IMCON', 'RELIDO', 'PROPIN', 'FISA', 'SCI',
        'CLASSIFIED BY', 'DERIVED FROM', 'DECLASSIFY ON',
        # Legacy CUI markings
        'FOUO', 'FOR OFFICIAL USE ONLY',
        'SBU', 'SENSITIVE BUT UNCLASSIFIED',
        'LES', 'LAW ENFORCEMENT SENSITIVE',
        'LIMDIS', 'LIMITED DISTRIBUTION',
    ]

    def __init__(self, threshold: int = 85):
        """
        Initialize fuzzy matcher
        threshold: Minimum similarity score (0-100) to consider a match
        """
        self.threshold = threshold
        self.enabled = HAS_RAPIDFUZZ

    def find_fuzzy_matches(self, text: str) -> List[Tuple[str, str, int, int]]:
        """
        Find fuzzy matches in text
        Returns: List of (matched_text, classification_term, score, position)
        """
        if not self.enabled:
            return []

        matches = []
        words = re.findall(r'\b\w+(?:\s+\w+){0,3}\b', text)

        for word_phrase in words:
            # Try fuzzy matching against known terms
            result = process.extractOne(
                word_phrase.upper(),
                self.CLASSIFICATION_TERMS,
                scorer=fuzz.ratio
            )

            if result and result[1] >= self.threshold:
                matched_term = result[0]
                score = result[1]

                # Find position in text
                position = text.upper().find(word_phrase.upper())
                if position != -1:
                    matches.append((word_phrase, matched_term, score, position))

        return matches


class LLMVerifier:
    """Verify classification markings using local LLM via Ollama"""

    def __init__(self, model: str = "qwen3:8b", enabled: bool = True):
        self.model = model
        self.enabled = enabled and HAS_OLLAMA

        if self.enabled:
            # Test Ollama connection
            try:
                ollama.list()
            except Exception as e:
                print(f"Warning: Ollama not available: {e}", file=sys.stderr)
                self.enabled = False

    def verify(self, text_sample: str, matches: List[Match]) -> Tuple[bool, float, str]:
        """
        Verify if matches are official classification markings
        Returns: (is_official, confidence, reasoning)
        """
        if not self.enabled or not matches:
            return True, 0.8, "LLM verification disabled"

        # Prepare match summary
        match_texts = list(set([m.text for m in matches[:10]]))

        prompt = f"""You are a US Government security classification expert. Analyze the following text excerpt for OFFICIAL classification markings.

POTENTIAL MARKINGS FOUND:
{chr(10).join(f'- "{m}"' for m in match_texts)}

TEXT EXCERPT (first 1500 chars):
{text_sample[:1500]}

TASK: Determine if these are OFFICIAL US Government classification markings (document headers, banners, portion marks, classification blocks per Executive Order 13526 and 32 CFR Part 2001) OR casual/colloquial usage of words like "secret" or "classified".

RESPOND EXACTLY IN THIS FORMAT:
VERDICT: [YES/NO]
CONFIDENCE: [0-100]
REASONING: [One sentence explanation]

Examples of OFFICIAL markings:
- "TOP SECRET//NOFORN" as a document header
- "(S)" at the start of a paragraph (portion marking)
- "CLASSIFIED BY: John Doe, Reason: 1.4(c), Declassify On: 20301231"

Examples of CASUAL usage (NOT official markings):
- "the secret to success"
- "this data can be classified as sensitive"
- "confidential conversation"

NOW ANALYZE:"""

        try:
            response = ollama.chat(
                model=self.model,
                messages=[{'role': 'user', 'content': prompt}],
                options={'temperature': 0.1}  # Low temperature for consistency
            )

            answer = response['message']['content'].strip()

            # Parse response
            verdict_match = re.search(r'VERDICT:\s*(YES|NO)', answer, re.IGNORECASE)
            confidence_match = re.search(r'CONFIDENCE:\s*(\d+)', answer)
            reasoning_match = re.search(r'REASONING:\s*(.+)', answer, re.DOTALL)

            is_official = verdict_match.group(1).upper() == 'YES' if verdict_match else True
            confidence = float(confidence_match.group(1)) / 100 if confidence_match else 0.7
            reasoning = reasoning_match.group(1).strip() if reasoning_match else "No reasoning provided"

            return is_official, confidence, reasoning

        except Exception as e:
            print(f"LLM verification error: {e}", file=sys.stderr)
            return True, 0.6, f"Error: {str(e)}"


class ClassificationScanner:
    """Main scanner orchestrating all detection layers"""

    SUPPORTED_EXTENSIONS = {
        '.pdf', '.docx', '.doc', '.txt', '.text', '.md', '.log', '.json'
    }

    def __init__(self, config: Dict):
        self.config = config
        self.extractor = TextExtractor()
        self.context_analyzer = ContextAnalyzer()
        self.fuzzy_matcher = FuzzyMatcher(
            threshold=config.get('fuzzy_threshold', 85)
        ) if config.get('fuzzy_matching', False) else None
        self.llm_verifier = LLMVerifier(
            model=config.get('llm_model', 'qwen3:8b'),
            enabled=config.get('use_llm', False)
        )
        self.patterns = ClassificationPatterns.get_all_patterns()

    def find_repeated_markings(self, text_sections: List[str]) -> Dict[str, int]:
        """
        Analyze text sections (e.g., page headers/footers) to find repeated markings.
        Returns dict of {normalized_marking: occurrence_count}
        """
        from collections import Counter
        marking_counts = Counter()

        for section in text_sections:
            # Find classification terms in this section
            section_upper = section.upper()

            # Check for common classification terms
            classification_terms = [
                'TOP SECRET', 'SECRET', 'CONFIDENTIAL', 'CUI',
                'FOUO', 'FOR OFFICIAL USE ONLY', 'SBU', 'LES',
                'NOFORN', 'ORCON', 'RELIDO'
            ]

            for term in classification_terms:
                if term in section_upper:
                    marking_counts[term] += 1

        return dict(marking_counts)

    def get_line_number(self, text: str, position: int) -> int:
        """Get line number for a position in text"""
        return text[:position].count('\n') + 1

    def determine_location(self, text: str, position: int, extracted_data: Dict) -> str:
        """Determine if match is in header, footer, body, or metadata"""
        # Check metadata
        metadata_text = json.dumps(extracted_data.get('metadata', {}))
        if metadata_text and text[position:position+50] in metadata_text:
            return 'metadata'

        # Check headers
        if 'headers' in extracted_data:
            headers_text = '\n'.join(extracted_data['headers'])
            if headers_text and text[position:position+50] in headers_text:
                return 'header'

        # Check footers
        if 'footers' in extracted_data:
            footers_text = '\n'.join(extracted_data['footers'])
            if footers_text and text[position:position+50] in footers_text:
                return 'footer'

        return 'body'

    def scan_text(self, text: str, extracted_data: Dict) -> List[Match]:
        """Scan text for classification markings using all detection layers"""
        matches = []

        # Layer 1: Pattern matching
        for category, name, pattern in self.patterns:
            for match in pattern.finditer(text):
                match_text = match.group(0).strip()
                position = match.start()

                # Context analysis
                context_info = self.context_analyzer.analyze(text, position)

                # Calculate confidence
                confidence = 0.5

                # Boost for official context
                if context_info['has_official']:
                    confidence += 0.3

                # Penalty for false positive indicators
                if context_info['has_false_positive']:
                    confidence -= 0.4

                # Boost for structural indicators
                if context_info['at_line_start']:
                    confidence += 0.15
                if context_info['has_slashes']:
                    confidence += 0.15

                # Boost for specific categories
                if category == 'structural':
                    confidence += 0.2
                if category == 'control' and context_info['has_official']:
                    confidence += 0.25
                if category == 'declassification':
                    # Declassification indicators are explicit evidence document was classified
                    confidence += 0.25

                # Determine location and boost for header/footer (classification banners)
                location = self.determine_location(text, position, extracted_data)
                if location in ['header', 'footer']:
                    # Classification banners typically appear in headers/footers
                    confidence += 0.2
                elif location == 'metadata':
                    # Metadata can contain official classification info
                    confidence += 0.1
                elif location == 'body' and not context_info['has_official']:
                    # Body text without official context is more likely false positive
                    confidence -= 0.2

                # Clamp confidence
                confidence = max(0.0, min(1.0, confidence))

                # Only keep matches above threshold
                threshold = self.config.get('confidence_threshold', 0.3)
                if confidence >= threshold:
                    matches.append(Match(
                        text=match_text,
                        position=position,
                        line_number=self.get_line_number(text, position),
                        confidence=confidence,
                        context=context_info['context'][:200],
                        match_type='pattern',
                        location=location
                    ))

        # Layer 2: Fuzzy matching (for variations and typos)
        if self.fuzzy_matcher and self.fuzzy_matcher.enabled:
            fuzzy_matches = self.fuzzy_matcher.find_fuzzy_matches(text)

            for word_phrase, term, score, position in fuzzy_matches:
                # Avoid duplicates from pattern matching
                if any(abs(m.position - position) < 10 for m in matches):
                    continue

                context_info = self.context_analyzer.analyze(text, position)

                # Fuzzy matches start with lower base confidence
                confidence = (score / 100) * 0.7

                if context_info['has_official']:
                    confidence += 0.2
                if context_info['has_false_positive']:
                    confidence -= 0.3

                # Determine location and boost for header/footer
                location = self.determine_location(text, position, extracted_data)
                if location in ['header', 'footer']:
                    confidence += 0.15  # Slightly lower boost for fuzzy matches
                elif location == 'metadata':
                    confidence += 0.1
                elif location == 'body' and not context_info['has_official']:
                    # Body text without official context is more likely false positive
                    confidence -= 0.2

                confidence = max(0.0, min(1.0, confidence))

                threshold = self.config.get('confidence_threshold', 0.3)
                if confidence >= threshold:
                    matches.append(Match(
                        text=f"{word_phrase} (fuzzy: {term})",
                        position=position,
                        line_number=self.get_line_number(text, position),
                        confidence=confidence,
                        context=context_info['context'][:200],
                        match_type='fuzzy',
                        location=location
                    ))

        # Remove duplicates (keep highest confidence)
        unique_matches = {}
        for match in matches:
            key = (match.position // 10) * 10  # Group nearby positions
            if key not in unique_matches or match.confidence > unique_matches[key].confidence:
                unique_matches[key] = match

        return sorted(unique_matches.values(), key=lambda m: m.confidence, reverse=True)

    def determine_classification_level(self, matches: List[Match]) -> ClassificationLevel:
        """Determine highest classification level from matches"""
        if not matches:
            return ClassificationLevel.UNCLASSIFIED

        # Check for classification levels in match texts
        all_text = ' '.join([m.text.upper() for m in matches])

        if 'TOP SECRET' in all_text or 'TOP-SECRET' in all_text:
            return ClassificationLevel.TOP_SECRET
        elif 'SECRET' in all_text:
            # Make sure it's not just "TOP SECRET"
            if 'TOP' not in all_text or all_text.count('SECRET') > all_text.count('TOP'):
                return ClassificationLevel.SECRET

        if 'CONFIDENTIAL' in all_text:
            return ClassificationLevel.CONFIDENTIAL
        if 'CUI' in all_text or 'CONTROLLED UNCLASSIFIED' in all_text:
            return ClassificationLevel.CUI
        # Check for legacy CUI markings (FOUO, SBU, LES, LIMDIS)
        if any(term in all_text for term in ['FOUO', 'FOR OFFICIAL USE ONLY', 'SBU',
                                               'SENSITIVE BUT UNCLASSIFIED', 'LES',
                                               'LAW ENFORCEMENT SENSITIVE', 'LIMDIS',
                                               'LIMITED DISTRIBUTION']):
            return ClassificationLevel.CUI

        return ClassificationLevel.UNKNOWN

    def scan_file(self, file_path: str) -> DetectionResult:
        """Scan a single file for classification markings"""
        import time
        start_time = time.time()

        file_size = os.path.getsize(file_path)
        file_hash = self.extractor.calculate_hash(file_path)

        # Extract text
        extracted_data, ocr_used = self.extractor.extract(file_path)

        # Combine all text sections
        text_parts = []

        if 'headers' in extracted_data:
            text_parts.extend(extracted_data['headers'])
        if 'body' in extracted_data:
            text_parts.extend(extracted_data['body'])
        if 'footers' in extracted_data:
            text_parts.extend(extracted_data['footers'])
        if 'metadata' in extracted_data:
            text_parts.append(json.dumps(extracted_data['metadata']))

        full_text = '\n\n'.join(str(part) for part in text_parts if part)

        if not full_text.strip():
            return DetectionResult(
                file_path=file_path,
                file_hash=file_hash,
                has_markings=False,
                classification_level=ClassificationLevel.UNCLASSIFIED,
                matches=[],
                overall_confidence=0.0,
                ocr_used=ocr_used,
                llm_verified=False,
                processing_time=time.time() - start_time,
                file_size=file_size
            )

        # Scan for matches
        matches = self.scan_text(full_text, extracted_data)

        # Page-association analysis: Boost confidence for repeated markings
        if matches and ('headers' in extracted_data or 'footers' in extracted_data):
            # Analyze headers and footers for repeated markings
            header_sections = extracted_data.get('headers', [])
            footer_sections = extracted_data.get('footers', [])
            all_sections = header_sections + footer_sections

            if len(all_sections) > 1:  # Only analyze if multiple pages
                repeated_markings = self.find_repeated_markings(all_sections)

                # Boost confidence for matches that appear repeatedly
                for match in matches:
                    match_text_upper = match.text.upper()

                    for repeated_term, count in repeated_markings.items():
                        if repeated_term in match_text_upper:
                            # Markings repeated on multiple pages are likely official banners
                            # Boost confidence based on repetition frequency
                            if count >= len(all_sections) * 0.5:  # On 50%+ of pages
                                match.confidence = min(1.0, match.confidence + 0.2)
                            elif count >= 3:  # On 3+ pages
                                match.confidence = min(1.0, match.confidence + 0.15)

        # LLM verification
        llm_verified = False
        if matches and self.llm_verifier.enabled:
            is_official, llm_confidence, reasoning = self.llm_verifier.verify(
                full_text, matches
            )
            llm_verified = True

            if not is_official:
                # LLM says these aren't official markings
                matches = []
            else:
                # Boost confidence based on LLM
                for match in matches:
                    match.confidence = max(match.confidence, llm_confidence)

        # Determine results
        has_markings = len(matches) > 0
        classification_level = self.determine_classification_level(matches)
        overall_confidence = max([m.confidence for m in matches]) if matches else 0.0

        processing_time = time.time() - start_time

        return DetectionResult(
            file_path=file_path,
            file_hash=file_hash,
            has_markings=has_markings,
            classification_level=classification_level,
            matches=matches,
            overall_confidence=overall_confidence,
            ocr_used=ocr_used,
            llm_verified=llm_verified,
            processing_time=processing_time,
            file_size=file_size
        )

    def scan_directory(self, directory: str, recursive: bool = True) -> List[DetectionResult]:
        """Scan all supported files in a directory"""
        path = Path(directory)
        files = []

        if recursive:
            for ext in self.SUPPORTED_EXTENSIONS:
                files.extend(path.rglob(f'*{ext}'))
        else:
            for ext in self.SUPPORTED_EXTENSIONS:
                files.extend(path.glob(f'*{ext}'))

        results = []
        workers = self.config.get('workers', 1)

        if workers > 1:
            # Parallel processing
            with ProcessPoolExecutor(max_workers=workers) as executor:
                futures = {executor.submit(self.scan_file, str(f)): f for f in files}

                if HAS_TQDM:
                    futures_iter = tqdm(
                        as_completed(futures),
                        total=len(futures),
                        desc="Scanning files"
                    )
                else:
                    futures_iter = as_completed(futures)

                for future in futures_iter:
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        file_path = futures[future]
                        print(f"Error scanning {file_path}: {e}", file=sys.stderr)
        else:
            # Sequential processing
            file_iter = tqdm(files, desc="Scanning files") if HAS_TQDM else files

            for file_path in file_iter:
                try:
                    result = self.scan_file(str(file_path))
                    results.append(result)
                except Exception as e:
                    print(f"Error scanning {file_path}: {e}", file=sys.stderr)

        return results


class ReportGenerator:
    """Generate various report formats"""

    @staticmethod
    def print_console_report(results: List[DetectionResult], show_all: bool = False):
        """Print formatted console report"""
        flagged = [r for r in results if r.has_markings]
        clean = [r for r in results if not r.has_markings]

        print("\n" + "=" * 80)
        print("CLASSIFICATION MARKING SCAN RESULTS")
        print("=" * 80)

        if flagged:
            print(f"\n⚠️  FLAGGED FILES: {len(flagged)}")
            print("-" * 80)

            for result in sorted(flagged, key=lambda x: x.classification_level):
                print(f"\n📄 File: {result.file_path}")
                print(f"   Classification: {result.classification_level.value}")
                print(f"   Confidence: {result.overall_confidence:.1%}")
                print(f"   Matches: {len(result.matches)}")
                print(f"   Hash: {result.file_hash[:16]}...")
                print(f"   OCR Used: {'Yes' if result.ocr_used else 'No'}")
                print(f"   LLM Verified: {'Yes' if result.llm_verified else 'No'}")

                # Show top matches
                for i, match in enumerate(result.matches[:3], 1):
                    print(f"\n   Match {i}:")
                    print(f"     Text: {match.text[:80]}")
                    print(f"     Line: {match.line_number}")
                    print(f"     Location: {match.location}")
                    print(f"     Type: {match.match_type}")
                    print(f"     Confidence: {match.confidence:.1%}")

                if len(result.matches) > 3:
                    print(f"\n   ... and {len(result.matches) - 3} more matches")

        else:
            print("\n✅ No classified documents found")

        if show_all and clean:
            print(f"\n\n✅ CLEAN FILES: {len(clean)}")
            print("-" * 80)
            for result in clean:
                print(f"   {result.file_path}")
        else:
            print(f"\n\n✅ Clean files: {len(clean)}")

        # Summary statistics
        print("\n" + "=" * 80)
        print("SUMMARY")
        print("-" * 80)
        print(f"Total files scanned: {len(results)}")
        print(f"Files flagged: {len(flagged)}")
        print(f"Clean files: {len(clean)}")

        if flagged:
            print("\nClassification breakdown:")
            level_counts = defaultdict(int)
            for r in flagged:
                level_counts[r.classification_level.value] += 1
            for level, count in sorted(level_counts.items()):
                print(f"  {level}: {count}")

        total_time = sum(r.processing_time for r in results)
        print(f"\nTotal processing time: {total_time:.2f}s")
        print("=" * 80)

    @staticmethod
    def generate_json_report(results: List[DetectionResult]) -> Dict:
        """Generate JSON report"""
        flagged = [r for r in results if r.has_markings]

        return {
            'scan_summary': {
                'total_files': len(results),
                'flagged_files': len(flagged),
                'clean_files': len(results) - len(flagged),
                'total_processing_time': sum(r.processing_time for r in results),
            },
            'flagged_files': [
                {
                    'file_path': r.file_path,
                    'file_hash': r.file_hash,
                    'file_size': r.file_size,
                    'classification_level': r.classification_level.value,
                    'overall_confidence': round(r.overall_confidence, 3),
                    'ocr_used': r.ocr_used,
                    'llm_verified': r.llm_verified,
                    'processing_time': round(r.processing_time, 3),
                    'matches': [
                        {
                            'text': m.text,
                            'line_number': m.line_number,
                            'confidence': round(m.confidence, 3),
                            'match_type': m.match_type,
                            'location': m.location,
                            'context': m.context[:100],
                        }
                        for m in r.matches
                    ]
                }
                for r in flagged
            ]
        }

    @staticmethod
    def generate_csv_report(results: List[DetectionResult]) -> str:
        """Generate CSV report"""
        import csv
        from io import StringIO

        output = StringIO()
        writer = csv.writer(output)

        # Header
        writer.writerow([
            'File Path', 'Classification Level', 'Confidence',
            'Match Count', 'File Hash', 'OCR Used', 'LLM Verified',
            'File Size', 'Processing Time'
        ])

        # Data
        for r in results:
            if r.has_markings:
                writer.writerow([
                    r.file_path,
                    r.classification_level.value,
                    f"{r.overall_confidence:.2%}",
                    len(r.matches),
                    r.file_hash,
                    r.ocr_used,
                    r.llm_verified,
                    r.file_size,
                    f"{r.processing_time:.2f}s"
                ])

        return output.getvalue()


def main():
    parser = argparse.ArgumentParser(
        description='Scan documents for US Government security classification markings (Desktop Edition)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Basic scan
  %(prog)s /path/to/documents

  # High recall mode with all features
  %(prog)s /path/to/docs --sensitivity high --use-llm --fuzzy-matching

  # Fast scan with parallel processing
  %(prog)s /path/to/docs --workers 4

  # Generate detailed reports
  %(prog)s /path/to/docs --output report.json --csv results.csv

  # Quarantine flagged files
  %(prog)s /path/to/docs --quarantine ./flagged_files

Supported formats:
  PDF (.pdf), Word (.docx, .doc), Text (.txt, .md, .log, .csv), JSON (.json)

Dependencies:
  Required: PyPDF2, python-docx
  Optional: pdfplumber, docx2python, rapidfuzz, ollama, tqdm

  Install: pip install -r requirements.txt

For LLM verification (optional):
  pip install ollama
  ollama pull qwen3:8b
        """
    )

    parser.add_argument('directory', help='Directory containing documents to scan')
    parser.add_argument('-r', '--recursive', action='store_true', default=True,
                       help='Scan subdirectories recursively (default: True)')
    parser.add_argument('--no-recursive', action='store_true',
                       help='Do not scan subdirectories')

    # Detection options
    parser.add_argument('--sensitivity', choices=['low', 'medium', 'high'], default='high',
                       help='Detection sensitivity (default: high)')
    parser.add_argument('--fuzzy-matching', action='store_true',
                       help='Enable fuzzy matching for variations and typos')
    parser.add_argument('--fuzzy-threshold', type=int, default=85,
                       help='Fuzzy matching threshold 0-100 (default: 85)')
    parser.add_argument('--confidence-threshold', type=float, default=0.3,
                       help='Minimum confidence to report (0.0-1.0, default: 0.3)')

    # LLM options
    parser.add_argument('--use-llm', action='store_true',
                       help='Enable LLM verification via Ollama')
    parser.add_argument('--llm-model', default='qwen3:8b',
                       help='Ollama model to use (default: qwen3:8b)')
    parser.add_argument('--no-llm', action='store_true',
                       help='Disable LLM verification')

    # Performance options
    parser.add_argument('--workers', type=int, default=1,
                       help='Number of parallel workers (default: 1)')

    # Output options
    parser.add_argument('--show-all', action='store_true',
                       help='Show all files including clean ones')
    parser.add_argument('-o', '--output', help='Output JSON report to file')
    parser.add_argument('--csv', help='Output CSV report to file')
    parser.add_argument('--quarantine', help='Copy flagged files to this directory')

    args = parser.parse_args()

    # Validate directory
    if not os.path.isdir(args.directory):
        print(f"Error: {args.directory} is not a valid directory", file=sys.stderr)
        sys.exit(1)

    # Check dependencies
    if not HAS_PYPDF2 and not HAS_PDFPLUMBER:
        print("Warning: Neither PyPDF2 nor pdfplumber installed. PDF files will be skipped.")
        print("Install with: pip install PyPDF2 pdfplumber")

    if not HAS_DOCX:
        print("Warning: python-docx not installed. DOCX files will be skipped.")
        print("Install with: pip install python-docx")

    if args.fuzzy_matching and not HAS_RAPIDFUZZ:
        print("Warning: rapidfuzz not installed. Fuzzy matching disabled.")
        print("Install with: pip install rapidfuzz")
        args.fuzzy_matching = False

    if args.use_llm and not args.no_llm and not HAS_OLLAMA:
        print("Error: Ollama not available but --use-llm specified", file=sys.stderr)
        print("Install with: pip install ollama", file=sys.stderr)
        print(f"Then run: ollama pull {args.llm_model}", file=sys.stderr)
        sys.exit(1)

    # Sensitivity presets
    if args.sensitivity == 'high':
        args.confidence_threshold = 0.3
        if not args.fuzzy_matching:
            args.fuzzy_matching = True
    elif args.sensitivity == 'medium':
        args.confidence_threshold = 0.5
    elif args.sensitivity == 'low':
        args.confidence_threshold = 0.7

    # Build configuration
    config = {
        'fuzzy_matching': args.fuzzy_matching,
        'fuzzy_threshold': args.fuzzy_threshold,
        'use_llm': args.use_llm and not args.no_llm,
        'llm_model': args.llm_model,
        'confidence_threshold': args.confidence_threshold,
        'workers': args.workers,
    }

    # Print configuration
    print("\nClassification Marking Scanner - Desktop Edition")
    print("=" * 80)
    print(f"Directory: {args.directory}")
    print(f"Recursive: {args.recursive and not args.no_recursive}")
    print(f"Sensitivity: {args.sensitivity}")
    print(f"Confidence Threshold: {config['confidence_threshold']:.1%}")
    print(f"Fuzzy Matching: {config['fuzzy_matching']}")
    print(f"LLM Verification: {config['use_llm']}")
    if config['use_llm']:
        print(f"LLM Model: {config['llm_model']}")
    print(f"Parallel Workers: {config['workers']}")
    print("=" * 80)

    # Create scanner
    scanner = ClassificationScanner(config)

    # Scan directory
    results = scanner.scan_directory(
        args.directory,
        recursive=(args.recursive and not args.no_recursive)
    )

    # Generate console report
    ReportGenerator.print_console_report(results, show_all=args.show_all)

    # Save JSON report
    if args.output:
        report_data = ReportGenerator.generate_json_report(results)
        with open(args.output, 'w') as f:
            json.dump(report_data, f, indent=2)
        print(f"\n📄 JSON report saved to: {args.output}")

    # Save CSV report
    if args.csv:
        csv_data = ReportGenerator.generate_csv_report(results)
        with open(args.csv, 'w') as f:
            f.write(csv_data)
        print(f"📄 CSV report saved to: {args.csv}")

    # Quarantine flagged files
    if args.quarantine:
        flagged = [r for r in results if r.has_markings]
        if flagged:
            os.makedirs(args.quarantine, exist_ok=True)
            for result in flagged:
                src = result.file_path
                dst = os.path.join(args.quarantine, os.path.basename(src))
                # Handle name conflicts
                counter = 1
                base_dst = dst
                while os.path.exists(dst):
                    name, ext = os.path.splitext(base_dst)
                    dst = f"{name}_{counter}{ext}"
                    counter += 1
                shutil.copy2(src, dst)
            print(f"\n📁 Quarantined {len(flagged)} files to: {args.quarantine}")

    # Exit with error code if classified documents found
    flagged_count = sum(1 for r in results if r.has_markings)
    sys.exit(1 if flagged_count > 0 else 0)


if __name__ == '__main__':
    main()
